from docx import Document
from datetime import datetime, timedelta
import os
import ctypes

def replace_text_in_paragraphs(paragraphs, old_text, new_text):
    for para in paragraphs:
        if old_text in para.text:
            para.text = para.text.replace(old_text, new_text)

def replace_text_in_tables(tables, old_text, new_text):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, old_text, new_text)

try:
    # 파일 경로 지정 (슬래시로 변경)
    original_file_path = 'C:/Users/ssaun_yr1jmer/OneDrive - Harlow International/0.2024년 업무폴더/1.일일업무/4.BOS/6월/example.docx'

    # 현재 날짜와 내일 날짜 구하기
    today = datetime.today()
    next_day = today + timedelta(days=1)

    # 날짜 형식을 영문식으로 변환
    today_str = today.strftime('%d %b %Y')
    next_day_str = next_day.strftime('%d %b %Y')

    # 메시지 박스 생성
    message = f"오늘 날짜: {today_str}\n문서 날짜: {next_day_str}\n실행할까요?"
    result = ctypes.windll.user32.MessageBoxW(0, message, "확인", 1)

    # 'Yes' 버튼 (ID: 6) 확인
    if result == 6:
        # 내일 날짜 이름으로 파일 복사
        new_file_name = f'example_{next_day_str}.docx'
        new_file_path = os.path.join('C:/Users/ssaun_yr1jmer/OneDrive - Harlow International/0.2024년 업무폴더/1.일일업무/4.BOS/6월/', new_file_name)

        # 원본 파일을 복사하여 새로운 파일 생성
        doc = Document(original_file_path)
        doc.save(new_file_path)
        print(f"파일이 성공적으로 생성되었습니다: {new_file_path}")

        # 생성된 파일 열기
        doc = Document(new_file_path)

        # 문서의 모든 단락에서 날짜 변경
        replace_text_in_paragraphs(doc.paragraphs, today_str, next_day_str)

        # 문서의 모든 테이블에서 날짜 변경
        replace_text_in_tables(doc.tables, today_str, next_day_str)

        # 변경된 내용 저장
        doc.save(new_file_path)
        print(f"문서 내 날짜가 성공적으로 변경되었습니다: {new_file_path}")

except Exception as e:
    print(f"오류가 발생했습니다: {e}")
    input("Press Enter to exit...")

